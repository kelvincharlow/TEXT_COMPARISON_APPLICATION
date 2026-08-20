from __future__ import annotations

import tempfile
import unittest
import zipfile
from pathlib import Path

from docx import Document

from backend.app.comparison.semantic_comparator import compare_semantic_changes
from backend.app.comparison.visual_redline import generate_visual_redline


def write_document(path: Path, paragraphs: list[str]) -> None:
    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    document.save(path)


class VisualRedlineTests(unittest.TestCase):
    def test_visual_redline_renders_insertions_deletions_and_modifications(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            original = root / "original.docx"
            revised = root / "revised.docx"
            output = root / "visual.docx"
            write_document(
                original,
                ["Stable opening", "Remove me", "Stable anchor", "Value 14", "Stable closing"],
            )
            write_document(
                revised,
                ["Stable opening", "Stable anchor", "Value 30", "Inserted note", "Stable closing"],
            )
            changes = compare_semantic_changes(original, revised)
            generate_visual_redline(original, revised, changes, output)

            Document(output)  # The generated package is readable by python-docx.
            with zipfile.ZipFile(output) as package:
                document_xml = package.read("word/document.xml")

        self.assertIn(b"Remove me", document_xml)
        self.assertIn(b"Inserted note", document_xml)
        self.assertIn(b'w:val="green"', document_xml)
        self.assertIn(b'w:val="FF0000"', document_xml)
        self.assertIn(b"POSTBANK DOCUMENT", document_xml)
        self.assertIn(b"VISUAL REDLINE COMPARISON", document_xml)
        self.assertIn(b"UNCHANGED", document_xml)
        self.assertIn(b"DELETED", document_xml)
        self.assertIn(b"INSERTED", document_xml)
        self.assertIn(b"<w:strike", document_xml)
        self.assertNotRegex(document_xml, br"<w:(?:ins|del)(?:\s|>)")


if __name__ == "__main__":
    unittest.main()
