"""Generate non-confidential DOCX pairs for the initial comparison experiment."""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

ROOT = Path(__file__).resolve().parent
SAMPLES = ROOT / "samples"


def add_heading(document: Document, reference: str, date_text: str) -> None:
    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run("POSTBANK DOCUMENT COMPARISON TEST LETTER")
    run.bold = True
    run.font.size = Pt(13)
    document.add_paragraph(f"Reference: {reference}")
    document.add_paragraph(f"Date: {date_text}")


def save_simple_pair() -> None:
    original = Document()
    add_heading(original, "PB/TEST/001", "19 August 2026")
    original.add_paragraph("Dear Sir or Madam,")
    original.add_paragraph(
        "Please provide the requested confirmation within 14 days of receiving this letter."
    )
    original.add_paragraph("Kind regards,\nTest Officer")
    original.save(SAMPLES / "simple_original.docx")

    revised = Document()
    add_heading(revised, "PB/TEST/001", "20 August 2026")
    revised.add_paragraph("Dear Sir or Madam,")
    revised.add_paragraph(
        "Please provide the requested written confirmation within 30 days of receiving this letter."
    )
    revised.add_paragraph("Please quote the reference above in your response.")
    revised.add_paragraph("Kind regards,\nTest Officer")
    revised.save(SAMPLES / "simple_revised.docx")


def save_table_pair() -> None:
    original = Document()
    add_heading(original, "PB/TEST/002", "19 August 2026")
    original.add_paragraph("The following test items require confirmation:")
    table = original.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for cell, value in zip(table.rows[0].cells, ("Item", "Reference", "Status")):
        cell.text = value
    for values in (("Account review", "REF-100", "Pending"), ("Address check", "REF-200", "Pending")):
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            cell.text = value
    original.save(SAMPLES / "table_original.docx")

    revised = Document()
    add_heading(revised, "PB/TEST/002", "19 August 2026")
    revised.add_paragraph("The following test items require confirmation:")
    table = revised.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for cell, value in zip(table.rows[0].cells, ("Item", "Reference", "Status")):
        cell.text = value
    for values in (
        ("Account review", "REF-101", "Complete"),
        ("Address check", "REF-200", "Pending"),
        ("Signature check", "REF-300", "Pending"),
    ):
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            cell.text = value
    revised.save(SAMPLES / "table_revised.docx")


def write_manifest() -> None:
    cases = [
        {
            "id": "simple_letter",
            "original": "samples/simple_original.docx",
            "revised": "samples/simple_revised.docx",
            "expected_changes": [
                "Date changed from 19 August 2026 to 20 August 2026",
                "The word 'written' was inserted",
                "14 days changed to 30 days",
                "A new reference-instruction paragraph was inserted",
            ],
        },
        {
            "id": "table_letter",
            "original": "samples/table_original.docx",
            "revised": "samples/table_revised.docx",
            "expected_changes": [
                "REF-100 changed to REF-101",
                "Account review status changed from Pending to Complete",
                "A Signature check table row was inserted",
            ],
        },
    ]
    (ROOT / "test_cases.json").write_text(json.dumps(cases, indent=2), encoding="utf-8")


def main() -> None:
    SAMPLES.mkdir(parents=True, exist_ok=True)
    save_simple_pair()
    save_table_pair()
    write_manifest()
    print(f"Created 2 synthetic document pairs in {SAMPLES}")


if __name__ == "__main__":
    main()

